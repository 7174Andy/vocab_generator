import numbers
import docx
import os
import random
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches

# Warning: Since this is a Korean private academy, the titles are in Korean. 

# contain method
def contain(list, num):
    for y in list:
        if y == num:
            return True

# get appropriate vocabulary list and bound
howManyVocab = int(input('Number of vocabs: '))
listNumber = int(input('Vocabulary List Number: '))
start = int(input('From: '))
end = int(input('to: '))
newWords = []
additionalWord = str(input('Additional words: '))
if additionalWord == '':
    pass
else:
    newWords.append(additionalWord)
    while additionalWord != '':
        additionalWord = str(input('Any more?: '))
        if additionalWord == '':
            break
        if contain(newWords, additionalWord):
            pass
        else: 
            newWords.append(additionalWord)

if howManyVocab == 120: 
    if listNumber == 1:
        document = docx.Document('칸 단어//A+ SAT Vocab (#1) 완성 2.docx')
    elif listNumber == 2:
        document = docx.Document('칸 단어//A+ SAT Vocab (#2) 완성.docx')
    elif listNumber == 3:
        document = docx.Document('칸 단어//A+ SAT Vocab (#3) 완성 2.docx')
    elif listNumber == 4:
        document = docx.Document('칸 단어//A+ SAT Vocab (#4) 완성.docx')
    elif listNumber == 5:
        document = docx.Document('칸 단어//A+ SAT Vocab (#5) 완성 2.docx')
    elif listNumber == 6:
        document = docx.Document('칸 단어//A+ SAT Vocab (#6) 완성.docx')
    elif listNumber == 7:
        document = docx.Document('칸 단어//A+ SAT Vocab (#7) 완성.docx')
    elif listNumber == 8:
        document = docx.Document('칸 단어//A+ SAT Vocab (#8) 완성.docx')
    elif listNumber == 9:
        document = docx.Document('칸 단어//A+ SAT Vocab (#9) 완성.docx')
    elif listNumber == 10:
        document = docx.Document('칸 단어//A+ SAT Vocab (#10) 완성.docx')
    elif listNumber == 11:
        document = docx.Document('칸 단어//A+ SAT Vocab (#11) 완성.docx')
    elif listNumber == 12:
        document = docx.Document('칸 단어//A+ SAT Vocab (#12) 완성.docx')
    elif listNumber == 13:
        document = docx.Document('칸 단어//A+ SAT Vocab (#13) 완성.docx')
    elif listNumber == 14:
        document = docx.Document('칸 단어//A+ SAT Vocab (#14) 완성.docx')
    elif listNumber == 15:
        document = docx.Document('칸 단어//A+ SAT Vocab (#15) 완성.docx')
    elif listNumber == 16:
        document = docx.Document('칸 단어//A+ SAT Vocab (#16) 완성.docx')
    elif listNumber == 17:
        document = docx.Document('칸 단어//A+ SAT Vocab (#17) 완성.docx')
    elif listNumber == 18:
        document = docx.Document('칸 단어//A+ SAT Vocab (#18) 완성.docx')
    else:
        print("List not found. \nPlease type the appropriate input.")
elif howManyVocab == 50:
    if listNumber == 1:
        document = docx.Document('칸 단어//50 SAT Vocab #1.docx')
    elif listNumber == 2:
        document = docx.Document('칸 단어//50 SAT Vocab #2.docx')
    elif listNumber == 3:
        document = docx.Document('칸 단어//50 SAT Vocab #3.docx')
    elif listNumber == 4:
        document = docx.Document('칸 단어//50 SAT Vocab #4.docx')
    elif listNumber == 5:
        document = docx.Document('칸 단어//50 SAT Vocab #5.docx')
    elif listNumber == 6:
        document = docx.Document('칸 단어//50 SAT Vocab #6.docx')
    elif listNumber == 7:
        document = docx.Document('칸 단어//50 SAT Vocab #7.docx')
    elif listNumber == 8:
        document = docx.Document('칸 단어//50 SAT Vocab #8.docx')
    else:
        print("List not found. \nPlease type the appropriate input.")
else:
    print("List not found. \nPlease type the appropriate input.")

# make three tests (original quiz + retest + re-retest)
n = 3
while n > 0:
    # pick random numbers
    x = 20
    numbers = []
    while x > 0:
        ran = random.randint(start,end)
        if contain(numbers, ran):
            while contain(numbers, ran):
                ran = random.randint(start,end)
        numbers.append(ran)
        x -= 1

    # create key file
    key = docx.Document()

    # title of the document
    header = key.sections[0].header
    if n == 3:
        header.paragraphs[0].text = '{} SAT Vocab {} Key (from {} to {})'.format(howManyVocab,listNumber, start, end)
    if n == 2: 
        header.paragraphs[0].text = '{} SAT Vocab {} 재시 Key (from {} to {})'.format(howManyVocab,listNumber, start, end)
    if n == 1:
        header.paragraphs[0].text = '{} SAT Vocab {} 재재시 Key (from {} to {})'.format(howManyVocab,listNumber, start, end)
    header.paragraphs[0].alignment = 1

    # add Kan logo
    logo = header.add_paragraph()
    logo_run = logo.add_run()
    logo_run.add_picture("KAN logo.jpg", width=Inches(0.25))
    logo.alignment = 1

    # write words
    for y in numbers:
        paragraph = key.add_paragraph('{}. {}'.format(numbers.index(y)+1, document.paragraphs[y - 1].text))
    if n == 3:
        for word in newWords:
            paragraph = key.add_paragraph('{}. {}'.format(newWords.index(word)+1+20, word))

    # save
    if n == 3:
        key.save('{} SAT Vocab {} Key (from {} to {}).docx'.format(howManyVocab,listNumber,start,end))
    if n == 2:
        key.save('{} SAT Vocab {} 재시 Key (from {} to {}).docx'.format(howManyVocab,listNumber,start,end))
    if n == 1:
        key.save('{} SAT Vocab {} 재재시 Key (from {} to {}).docx'.format(howManyVocab,listNumber,start,end))
    n -= 1

n = 3
while n > 0:
    # open appropriate key files
    if n == 3:
        key = docx.Document('{} SAT Vocab {} Key (from {} to {}).docx'.format(howManyVocab,listNumber,start,end))
    elif n == 2:
        key = docx.Document('{} SAT Vocab {} 재시 Key (from {} to {}).docx'.format(howManyVocab,listNumber,start,end))
    elif n == 1:
        key = docx.Document('{} SAT Vocab {} 재재시 Key (from {} to {}).docx'.format(howManyVocab,listNumber,start,end))

    # create new file for quiz
    quiz = docx.Document()

    # add header
    header = quiz.sections[0].header
    if n == 3:
        header.paragraphs[0].text = '{} SAT Vocab {} (from {} to {})'.format(howManyVocab,listNumber,start,end)
    if n == 2:
        header.paragraphs[0].text = '{} SAT Vocab {} 재시 (from {} to {})'.format(howManyVocab,listNumber,start,end)
    if n == 1:
        header.paragraphs[0].text = '{} SAT Vocab {} 재재시 (from {} to {})'.format(howManyVocab,listNumber,start,end)
    header.paragraphs[0].alignment = 1

    # add Kan logo
    logo = header.add_paragraph()
    logo_run = logo.add_run()
    logo_run.add_picture("KAN logo.jpg", width=Inches(0.25))
    logo.alignment = 1

    # add name + date + class
    header.add_paragraph('Name: \t\t\t\tClass: \t\t\t\tDate:')

    # add words
    for word in key.paragraphs:
        index = word.text.index(':')
        paragraph = quiz.add_paragraph(word.text[0:index+1])

    # save quiz
    if n == 3:
        quiz.save('{} SAT Vocab {} 단어시험 (from {} to {}).docx'.format(howManyVocab,listNumber,start,end))
    elif n == 2:
        quiz.save('{} SAT Vocab {} 재시 (from {} to {}).docx'.format(howManyVocab,listNumber,start,end))
    elif n == 1:
        quiz.save('{} SAT Vocab {} 재재시(from {} to {}).docx'.format(howManyVocab,listNumber,start,end))
    n -= 1
    
print("Complete!")